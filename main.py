#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""中药处方记录系统 - 打印机魂大悦"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime, timedelta
import sqlite3
import os
import win32print
import win32api
from collections import Counter
import re
import json

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告：未安装python-docx库")


# 处方文件保存文件夹
PRESCRIPTION_FOLDER = "处方记录"


class Settings:
    def __init__(self, settings_file="settings.json"):
        self.settings_file = settings_file
        self.default_doctor = ""
        self.default_phone = ""
        self.smart_completion_enabled = True
        
        # 诊所信息设置
        self.clinic_name = ""  # 诊所名称
        self.clinic_address = ""  # 诊所地址
        self.clinic_phone = ""  # 诊所电话
        self.clinic_license = ""  # 执业许可证号
        
        # 压缩参数（可调整）
        self.font_size = 9  # 基础字体大小
        self.line_spacing = 0.85  # 行间距系数（0.7-1.2）
        self.safety_margin = 1.5  # 安全系数（1.3-2.0）
        self.margin_size = 0.2  # 边距大小（cm）
        self.empty_ratio = 4.0  # 留空比例（默认4.0）
        
        self.load_settings()
    
    def load_settings(self):
        try:
            if os.path.exists(self.settings_file):
                with open(self.settings_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.default_doctor = data.get('default_doctor', '')
                    self.default_phone = data.get('default_phone', '')
                    self.smart_completion_enabled = data.get('smart_completion_enabled', True)
                    
                    # 加载诊所信息
                    self.clinic_name = data.get('clinic_name', '')
                    self.clinic_address = data.get('clinic_address', '')
                    self.clinic_phone = data.get('clinic_phone', '')
                    self.clinic_license = data.get('clinic_license', '')
                    
                    # 加载压缩参数
                    self.font_size = data.get('font_size', 9)
                    self.line_spacing = data.get('line_spacing', 0.85)
                    self.safety_margin = data.get('safety_margin', 1.5)
                    self.margin_size = data.get('margin_size', 0.2)
                    self.empty_ratio = float(data.get('empty_ratio', 4.0))
        except:
            pass
    
    def save_settings(self):
        try:
            with open(self.settings_file, 'w', encoding='utf-8') as f:
                json.dump({
                    'default_doctor': self.default_doctor,
                    'default_phone': self.default_phone,
                    'smart_completion_enabled': self.smart_completion_enabled,
                    
                    # 保存诊所信息
                    'clinic_name': self.clinic_name,
                    'clinic_address': self.clinic_address,
                    'clinic_phone': self.clinic_phone,
                    'clinic_license': self.clinic_license,
                    
                    # 保存压缩参数
                    'font_size': self.font_size,
                    'line_spacing': self.line_spacing,
                    'safety_margin': self.safety_margin,
                    'margin_size': self.margin_size,
                    'empty_ratio': self.empty_ratio
                }, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存设置失败：{e}")


class ContextMenu:
    def __init__(self, widget):
        self.widget = widget
        self.menu = tk.Menu(widget, tearoff=0, font=("Microsoft YaHei", 9))
        self.menu.add_command(label="撤销", command=self.undo, accelerator="Ctrl+Z")
        self.menu.add_separator()
        self.menu.add_command(label="剪切", command=self.cut, accelerator="Ctrl+X")
        self.menu.add_command(label="复制", command=self.copy, accelerator="Ctrl+C")
        self.menu.add_command(label="粘贴", command=self.paste, accelerator="Ctrl+V")
        self.menu.add_command(label="删除", command=self.delete, accelerator="Delete")
        self.menu.add_separator()
        self.menu.add_command(label="全选", command=self.select_all, accelerator="Ctrl+A")
        self.widget.bind("<Button-3>", self.show_menu)
    
    def show_menu(self, event):
        try:
            self.widget.focus_set()
            self.menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.menu.grab_release()
    
    def undo(self):
        try:
            if isinstance(self.widget, tk.Text):
                self.widget.edit_undo()
            else:
                self.widget.event_generate("<<Undo>>")
        except:
            pass
    
    def cut(self):
        try:
            self.widget.event_generate("<<Cut>>")
        except:
            pass
    
    def copy(self):
        try:
            self.widget.event_generate("<<Copy>>")
        except:
            pass
    
    def paste(self):
        try:
            self.widget.event_generate("<<Paste>>")
        except:
            pass
    
    def delete(self):
        try:
            if isinstance(self.widget, tk.Text):
                if self.widget.tag_ranges("sel"):
                    self.widget.delete("sel.first", "sel.last")
            else:
                self.widget.event_generate("<<Clear>>")
        except:
            pass
    
    def select_all(self):
        try:
            if isinstance(self.widget, tk.Text):
                self.widget.tag_add("sel", "1.0", "end")
                self.widget.mark_set("insert", "1.0")
            else:
                self.widget.select_range(0, "end")
                self.widget.icursor("end")
        except:
            pass


class SmartCompletionPanel:
    def __init__(self, parent, db_file, settings, on_select_callback=None):
        self.parent = parent
        self.db_file = db_file
        self.settings = settings
        self.on_select_callback = on_select_callback
        self.enabled = tk.BooleanVar(value=settings.smart_completion_enabled)
        self.all_words = []
        self.create_widgets()
        self.load_words_from_database()
    
    def create_widgets(self):
        self.main_frame = ttk.LabelFrame(self.parent, text="智能补全", padding="5")
        
        control_frame = ttk.Frame(self.main_frame)
        control_frame.pack(fill="x", pady=(0, 5))
        
        self.switch = ttk.Checkbutton(control_frame, text="启用智能补全",
            variable=self.enabled, command=self.toggle_completion)
        self.switch.pack(side="left", padx=5)
        
        ttk.Button(control_frame, text="刷新词库", command=self.load_words_from_database,
            width=10).pack(side="left", padx=5)
        
        self.stats_label = ttk.Label(control_frame, text="词库：0个词条")
        self.stats_label.pack(side="left", padx=10)
        
        self.notebook = ttk.Notebook(self.main_frame)
        self.notebook.pack(fill="both", expand=True)
        
        self.categories = {
            "常用药材": {"words": []},
            "常用诊断": {"words": []},
            "常用处方": {"words": []},
            "常用用法": {"words": []},
            "全部词条": {"words": []}
        }
        
        for cat_name in self.categories:
            frame = ttk.Frame(self.notebook, padding="5")
            self.notebook.add(frame, text=cat_name)
            self.categories[cat_name]["frame"] = frame
            
            search_frame = ttk.Frame(frame)
            search_frame.pack(fill="x", pady=(0, 5))
            ttk.Label(search_frame, text="筛选：").pack(side="left")
            search_entry = ttk.Entry(search_frame, width=30)
            search_entry.pack(side="left", padx=5)
            search_entry.bind('<KeyRelease>', lambda e, cat=cat_name: self.filter_words(cat, e))
            self.categories[cat_name]["search"] = search_entry
            
            canvas_frame = ttk.Frame(frame)
            canvas_frame.pack(fill="both", expand=True)
            canvas = tk.Canvas(canvas_frame, height=80, bg="white")
            scrollbar = ttk.Scrollbar(canvas_frame, orient="vertical", command=canvas.yview)
            scrollable_frame = ttk.Frame(canvas)
            scrollable_frame.bind("<Configure>", lambda e, c=canvas: c.configure(scrollregion=c.bbox("all")))
            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)
            scrollbar.pack(side="right", fill="y")
            canvas.pack(side="left", fill="both", expand=True)
            self.categories[cat_name]["canvas"] = canvas
            self.categories[cat_name]["scrollable"] = scrollable_frame
        
        ttk.Label(self.main_frame, text="点击词语可插入到当前输入框",
            foreground="gray").pack(pady=2)
    
    def toggle_completion(self):
        self.settings.smart_completion_enabled = self.enabled.get()
        self.settings.save_settings()
        if self.enabled.get():
            self.notebook.state(['!disabled'])
        else:
            self.notebook.state(['disabled'])
    
    def load_words_from_database(self):
        self.all_words = []
        try:
            conn = sqlite3.connect(self.db_file)
            cursor = conn.cursor()
            cursor.execute("SELECT patient_name, diagnosis, prescription, usage, doctor FROM prescriptions")
            rows = cursor.fetchall()
            conn.close()
            
            medicine_words = Counter()
            diagnosis_words = Counter()
            prescription_phrases = Counter()
            usage_words = Counter()
            
            for row in rows:
                name, diagnosis, prescription, usage, doctor = row
                if diagnosis:
                    words = re.split(r'[，,。、；：:；\s\[\]【】]+', diagnosis)
                    for w in words:
                        if len(w.strip()) >= 2:
                            diagnosis_words[w.strip()] += 1
                
                if prescription:
                    lines = prescription.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:
                            parts = re.split(r'[，,、\s]+', line)
                            for part in parts:
                                part = part.strip()
                                part = re.sub(r'\d+.*$', '', part)
                                part = re.sub(r'[克g粒片包钱两升].*$', '', part)
                                if 2 <= len(part) <= 6:
                                    medicine_words[part] += 1
                            if len(line) >= 4:
                                prescription_phrases[line] += 1
                if usage:
                    usage_words[usage] += 1
            
            self.categories["常用药材"]["words"] = [w for w, c in medicine_words.most_common(100)]
            self.categories["常用诊断"]["words"] = [w for w, c in diagnosis_words.most_common(50)]
            self.categories["常用处方"]["words"] = [w for w, c in prescription_phrases.most_common(30)]
            self.categories["常用用法"]["words"] = [w for w, c in usage_words.most_common(20)]
            all_words = list(medicine_words.keys()) + list(diagnosis_words.keys()) + list(usage_words.keys())
            self.categories["全部词条"]["words"] = sorted(set(all_words))
            self.all_words = self.categories["全部词条"]["words"]
            self.stats_label.config(text=f"词库：{len(self.all_words)}个词条")
            for cat_name in self.categories:
                self.display_words(cat_name)
        except Exception as e:
            print(f"加载词库失败：{e}")
    
    def display_words(self, category):
        scrollable = self.categories[category]["scrollable"]
        for widget in scrollable.winfo_children():
            widget.destroy()
        words = self.categories[category]["words"]
        if not words:
            ttk.Label(scrollable, text="暂无数据，保存处方后将自动学习",
                foreground="gray").pack(pady=10)
            return
        row_frame = None
        for i, word in enumerate(words):
            if i % 6 == 0:
                row_frame = ttk.Frame(scrollable)
                row_frame.pack(fill="x", pady=2)
            btn = tk.Button(row_frame, text=word, width=min(12, max(6, len(word) + 2)),
                command=lambda w=word: self.on_word_click(w), relief="raised", bd=1)
            btn.pack(side="left", padx=2, pady=1)
    
    def filter_words(self, category, event):
        search_text = self.categories[category]["search"].get().strip().lower()
        scrollable = self.categories[category]["scrollable"]
        for widget in scrollable.winfo_children():
            widget.destroy()
        all_words = self.categories[category]["words"]
        words = [w for w in all_words if search_text in w.lower()] if search_text else all_words
        if not words:
            ttk.Label(scrollable, text="未找到匹配词条", foreground="gray").pack(pady=10)
            return
        row_frame = None
        for i, word in enumerate(words[:50]):
            if i % 6 == 0:
                row_frame = ttk.Frame(scrollable)
                row_frame.pack(fill="x", pady=2)
            btn = tk.Button(row_frame, text=word, width=min(12, max(6, len(word) + 2)),
                command=lambda w=word: self.on_word_click(w), relief="raised", bd=1)
            btn.pack(side="left", padx=2, pady=1)
    
    def on_word_click(self, word):
        if self.on_select_callback:
            self.on_select_callback(word)
    
    def get_frame(self):
        return self.main_frame
    
    def is_enabled(self):
        return self.enabled.get()


class PrescriptionSystem:
    def __init__(self, root):
        self.root = root
        self.root.title("中药处方记录系统 - 中药忘开单")
        self.root.geometry("1400x900")
        self.settings = Settings()
        self.db_file = "prescriptions.db"
        self.template_file = "处方打印样本.docx"
        
        # 创建处方保存文件夹
        self.ensure_prescription_folder()
        
        self.init_database()
        self.create_widgets()
        self.refresh_printers()
    
    def ensure_prescription_folder(self):
        """确保处方文件夹存在"""
        if not os.path.exists(PRESCRIPTION_FOLDER):
            os.makedirs(PRESCRIPTION_FOLDER)
            print(f"已创建文件夹：{PRESCRIPTION_FOLDER}")
    
    def init_database(self):
        conn = sqlite3.connect(self.db_file)
        cursor = conn.cursor()
        cursor.execute("CREATE TABLE IF NOT EXISTS prescriptions (id INTEGER PRIMARY KEY AUTOINCREMENT, patient_name TEXT NOT NULL, gender TEXT, age TEXT, phone TEXT, diagnosis TEXT, prescription TEXT NOT NULL, usage TEXT, doctor TEXT, doctor_phone TEXT, create_time TEXT NOT NULL, print_time TEXT)")
        conn.commit()
        conn.close()
    
    def create_widgets(self):
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill="both", expand=True, padx=10, pady=10)
        input_frame = ttk.Frame(notebook)
        notebook.add(input_frame, text="处方录入")
        self.create_input_page(input_frame)
        query_frame = ttk.Frame(notebook)
        notebook.add(query_frame, text="历史查询")
        self.create_query_page(query_frame)
        settings_frame = ttk.Frame(notebook)
        notebook.add(settings_frame, text="设置")
        self.create_settings_page(settings_frame)
    
    def create_settings_page(self, parent):
        # 创建主框架，使用grid布局确保全页显示
        parent.grid_rowconfigure(0, weight=1)
        parent.grid_columnconfigure(0, weight=1)
        
        # 创建一个带滚动条的画布
        canvas = tk.Canvas(parent, bg="white")
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        main_frame = ttk.Frame(canvas, padding="20")
        
        main_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=main_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 使用grid布局确保全页显示
        canvas.grid(row=0, column=0, sticky="nsew")
        scrollbar.grid(row=0, column=1, sticky="ns")
        
        # 鼠标滚轮绑定
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        # 设置main_frame的宽度，确保内容居中显示
        def configure_frame(event=None):
            # 更新滚动区域
            canvas.configure(scrollregion=canvas.bbox("all"))
            # 设置画布宽度为窗口宽度
            canvas_width = canvas.winfo_width()
            if canvas_width > 1:  # 确保窗口已经初始化
                canvas.itemconfig(canvas.find_withtag("all")[0], width=canvas_width)
        
        canvas.bind("<Configure>", configure_frame)
        
        # 标题
        title_label = ttk.Label(main_frame, text="系统设置", font=("Microsoft YaHei", 16, "bold"))
        title_label.pack(pady=(0, 20), anchor="center")
        
        # 诊所信息设置
        clinic_frame = ttk.LabelFrame(main_frame, text="诊所信息", padding="15")
        clinic_frame.pack(fill="both", expand=True, pady=10)
        
        # 诊所名称
        clinic_name_row = ttk.Frame(clinic_frame)
        clinic_name_row.pack(fill="x", pady=5)
        ttk.Label(clinic_name_row, text="诊所名称：", width=12).pack(side="left")
        self.clinic_name_entry = ttk.Entry(clinic_name_row, width=40)
        self.clinic_name_entry.insert(0, self.settings.clinic_name)
        self.clinic_name_entry.pack(side="left", padx=10)
        
        # 诊所地址
        clinic_address_row = ttk.Frame(clinic_frame)
        clinic_address_row.pack(fill="x", pady=5)
        ttk.Label(clinic_address_row, text="诊所地址：", width=12).pack(side="left")
        self.clinic_address_entry = ttk.Entry(clinic_address_row, width=40)
        self.clinic_address_entry.insert(0, self.settings.clinic_address)
        self.clinic_address_entry.pack(side="left", padx=10)
        
        # 诊所电话
        clinic_phone_row = ttk.Frame(clinic_frame)
        clinic_phone_row.pack(fill="x", pady=5)
        ttk.Label(clinic_phone_row, text="诊所电话：", width=12).pack(side="left")
        self.clinic_phone_entry = ttk.Entry(clinic_phone_row, width=40)
        self.clinic_phone_entry.insert(0, self.settings.clinic_phone)
        self.clinic_phone_entry.pack(side="left", padx=10)
        
        # 执业许可证号
        clinic_license_row = ttk.Frame(clinic_frame)
        clinic_license_row.pack(fill="x", pady=5)
        ttk.Label(clinic_license_row, text="执业许可证号：", width=12).pack(side="left")
        self.clinic_license_entry = ttk.Entry(clinic_license_row, width=40)
        self.clinic_license_entry.insert(0, self.settings.clinic_license)
        self.clinic_license_entry.pack(side="left", padx=10)
        
        # 医生信息设置
        doctor_frame = ttk.LabelFrame(main_frame, text="默认医生信息", padding="15")
        doctor_frame.pack(fill="both", expand=True, pady=10)
        
        doctor_row1 = ttk.Frame(doctor_frame)
        doctor_row1.pack(fill="x", pady=5)
        ttk.Label(doctor_row1, text="医    生：", width=12).pack(side="left")
        self.doctor_entry = ttk.Entry(doctor_row1, width=40)
        self.doctor_entry.insert(0, self.settings.default_doctor)
        self.doctor_entry.pack(side="left", padx=10)
        
        doctor_row2 = ttk.Frame(doctor_frame)
        doctor_row2.pack(fill="x", pady=5)
        ttk.Label(doctor_row2, text="医生电话：", width=12).pack(side="left")
        self.phone_entry = ttk.Entry(doctor_row2, width=40)
        self.phone_entry.insert(0, self.settings.default_phone)
        self.phone_entry.pack(side="left", padx=10)
        
        # 压缩参数设置
        compress_frame = ttk.LabelFrame(main_frame, text="页面压缩参数（调整后保存）", padding="15")
        compress_frame.pack(fill="both", expand=True, pady=10)
        
        # 字体大小
        font_row = ttk.Frame(compress_frame)
        font_row.pack(fill="x", pady=5)
        ttk.Label(font_row, text="字体大小：", width=15).pack(side="left")
        self.font_size_var = tk.StringVar(value=str(self.settings.font_size))
        font_spinbox = ttk.Spinbox(font_row, from_=6, to=12, width=10, 
                                    textvariable=self.font_size_var)
        font_spinbox.pack(side="left", padx=10)
        ttk.Label(font_row, text="推荐: 7-9（越小越紧凑）", foreground="gray").pack(side="left")
        
        # 行间距
        spacing_row = ttk.Frame(compress_frame)
        spacing_row.pack(fill="x", pady=5)
        ttk.Label(spacing_row, text="行间距系数：", width=15).pack(side="left")
        self.line_spacing_var = tk.StringVar(value=str(self.settings.line_spacing))
        spacing_spinbox = ttk.Spinbox(spacing_row, from_=0.6, to=1.2, increment=0.05, 
                                       width=10, textvariable=self.line_spacing_var)
        spacing_spinbox.pack(side="left", padx=10)
        ttk.Label(spacing_row, text="推荐: 0.7-0.9（越小越紧凑）", foreground="gray").pack(side="left")
        
        # 安全系数（留空比例）
        safety_row = ttk.Frame(compress_frame)
        safety_row.pack(fill="x", pady=5)
        ttk.Label(safety_row, text="留空安全系数：", width=15).pack(side="left")
        self.safety_margin_var = tk.StringVar(value=str(self.settings.safety_margin))
        safety_spinbox = ttk.Spinbox(safety_row, from_=1.3, to=2.5, increment=0.1,
                                      width=10, textvariable=self.safety_margin_var)
        safety_spinbox.pack(side="left", padx=10)
        ttk.Label(safety_row, text="推荐: 1.5-2.0（越大留空越多）", foreground="gray").pack(side="left")
        
        # 边距大小
        margin_row = ttk.Frame(compress_frame)
        margin_row.pack(fill="x", pady=5)
        ttk.Label(margin_row, text="边距大小：", width=15).pack(side="left")
        self.margin_var = tk.StringVar(value=str(self.settings.margin_size))
        margin_spinbox = ttk.Spinbox(margin_row, from_=0.1, to=0.5, increment=0.05,
                                      width=10, textvariable=self.margin_var)
        margin_spinbox.pack(side="left", padx=10)
        ttk.Label(margin_row, text="推荐: 0.15-0.25（越小越紧凑）", foreground="gray").pack(side="left")
        
        # 留空比例
        empty_row = ttk.Frame(compress_frame)
        empty_row.pack(fill="x", pady=5)
        ttk.Label(empty_row, text="留空比例：", width=15).pack(side="left")
        self.empty_ratio_var = tk.StringVar(value=str(self.settings.empty_ratio))
        empty_spinbox = ttk.Spinbox(empty_row, from_=1.5, to=10.0, increment=0.1,
                                    width=10, textvariable=self.empty_ratio_var)
        empty_spinbox.pack(side="left", padx=10)
        ttk.Label(empty_row, text="推荐: 2.5-5.0（数值越大留空越多）", foreground="gray").pack(side="left")
        
        # 预设按钮
        preset_frame = ttk.Frame(compress_frame)
        preset_frame.pack(fill="x", pady=10)
        ttk.Label(preset_frame, text="快速预设：").pack(side="left", padx=5)
        ttk.Button(preset_frame, text="极简模式", command=self.preset_minimal, width=10).pack(side="left", padx=5)
        ttk.Button(preset_frame, text="标准模式", command=self.preset_standard, width=10).pack(side="left", padx=5)
        ttk.Button(preset_frame, text="宽松模式", command=self.preset_loose, width=10).pack(side="left", padx=5)
        
        # 智能文本设置
        smart_frame = ttk.LabelFrame(main_frame, text="智能文本设置", padding="15")
        smart_frame.pack(fill="both", expand=True, pady=10)
        self.smart_var = tk.BooleanVar(value=self.settings.smart_completion_enabled)
        ttk.Checkbutton(smart_frame, text="启用智能文本补全", variable=self.smart_var).pack(anchor="w", pady=5)
        
        # ========== 操作按钮框（放在智能文本设置下方）==========
        action_frame = ttk.LabelFrame(main_frame, text="操作", padding="15")
        action_frame.pack(fill="both", expand=True, pady=10)
        
        action_btn_frame = ttk.Frame(action_frame)
        action_btn_frame.pack(pady=10)
        
        # 使用tk.Button确保按钮可见
        save_btn = tk.Button(action_btn_frame, text="确认保存", command=self.save_settings, 
                            width=15, height=2, bg="#4CAF50", fg="white", font=("Microsoft YaHei", 10, "bold"))
        save_btn.pack(side="left", padx=20)
        
        cancel_btn = tk.Button(action_btn_frame, text="取消", command=self.cancel_settings, 
                              width=15, height=2, bg="#f44336", fg="white", font=("Microsoft YaHei", 10, "bold"))
        cancel_btn.pack(side="left", padx=20)
        
        # 模板说明
        template_frame = ttk.LabelFrame(main_frame, text="模板文件", padding="15")
        template_frame.pack(fill="both", expand=True, pady=10)
        ttk.Label(template_frame, text="使用58mm小票格式打印", foreground="gray").pack(anchor="w", pady=2)
        ttk.Label(template_frame, text="页面高度根据内容自动调整，强制单页", foreground="blue").pack(anchor="w", pady=2)
        
        # 文件保存设置
        folder_frame = ttk.LabelFrame(main_frame, text="文件保存设置", padding="15")
        folder_frame.pack(fill="both", expand=True, pady=10)
        ttk.Label(folder_frame, text=f"处方文件保存位置：{os.path.abspath(PRESCRIPTION_FOLDER)}", foreground="blue").pack(anchor="w", pady=2)
        ttk.Label(folder_frame, text="（程序会自动创建该文件夹并保存所有处方文件）", foreground="gray").pack(anchor="w", pady=2)
    
    def preset_minimal(self):
        """极简模式 - 最紧凑"""
        self.font_size_var.set("7")
        self.line_spacing_var.set("0.7")
        self.safety_margin_var.set("1.3")
        self.margin_var.set("0.15")
        self.empty_ratio_var.set("2.5")
    
    def preset_standard(self):
        """标准模式"""
        self.font_size_var.set("9")
        self.line_spacing_var.set("0.85")
        self.safety_margin_var.set("1.5")
        self.margin_var.set("0.2")
        self.empty_ratio_var.set("4.0")
    
    def preset_loose(self):
        """宽松模式"""
        self.font_size_var.set("10")
        self.line_spacing_var.set("1.0")
        self.safety_margin_var.set("1.3")
        self.margin_var.set("0.25")
        self.empty_ratio_var.set("5.5")
    
    def save_settings(self):
        self.settings.smart_completion_enabled = self.smart_var.get()
        self.settings.default_doctor = self.doctor_entry.get().strip()
        self.settings.default_phone = self.phone_entry.get().strip()
        
        # 保存诊所信息
        self.settings.clinic_name = self.clinic_name_entry.get().strip()
        self.settings.clinic_address = self.clinic_address_entry.get().strip()
        self.settings.clinic_phone = self.clinic_phone_entry.get().strip()
        self.settings.clinic_license = self.clinic_license_entry.get().strip()
        
        # 保存压缩参数
        try:
            self.settings.font_size = int(self.font_size_var.get())
            self.settings.line_spacing = float(self.line_spacing_var.get())
            self.settings.safety_margin = float(self.safety_margin_var.get())
            self.settings.margin_size = float(self.margin_var.get())
            self.settings.empty_ratio = float(self.empty_ratio_var.get())
        except:
            messagebox.showwarning("提示", "压缩参数格式错误！")
            return
        
        self.settings.save_settings()
        
        if hasattr(self, 'completion_panel'):
            if self.settings.smart_completion_enabled:
                self.completion_panel.enabled.set(True)
                self.completion_panel.notebook.state(['!disabled'])
            else:
                self.completion_panel.enabled.set(False)
                self.completion_panel.notebook.state(['disabled'])
        if hasattr(self, 'doctor_label'):
            if self.settings.default_doctor:
                self.doctor_label.config(text=self.settings.default_doctor)
            else:
                self.doctor_label.config(text="未设置（请在设置页面配置）")
        messagebox.showinfo("成功", "设置已保存！")
    
    def cancel_settings(self):
        self.smart_var.set(self.settings.smart_completion_enabled)
        self.doctor_entry.delete(0, tk.END)
        self.doctor_entry.insert(0, self.settings.default_doctor)
        self.phone_entry.delete(0, tk.END)
        self.phone_entry.insert(0, self.settings.default_phone)
        
        # 恢复诊所信息
        self.clinic_name_entry.delete(0, tk.END)
        self.clinic_name_entry.insert(0, self.settings.clinic_name)
        self.clinic_address_entry.delete(0, tk.END)
        self.clinic_address_entry.insert(0, self.settings.clinic_address)
        self.clinic_phone_entry.delete(0, tk.END)
        self.clinic_phone_entry.insert(0, self.settings.clinic_phone)
        self.clinic_license_entry.delete(0, tk.END)
        self.clinic_license_entry.insert(0, self.settings.clinic_license)
        
        self.font_size_var.set(str(self.settings.font_size))
        self.line_spacing_var.set(str(self.settings.line_spacing))
        self.safety_margin_var.set(str(self.settings.safety_margin))
        self.margin_var.set(str(self.settings.margin_size))
        self.empty_ratio_var.set(str(self.settings.empty_ratio))
    
    def create_input_page(self, parent):
        main_container = ttk.Frame(parent)
        main_container.pack(fill="both", expand=True)
        
        left_frame = ttk.Frame(main_container, padding="15")
        left_frame.pack(side="left", fill="both", expand=True)
        ttk.Label(left_frame, text="中药处方开具", font=("Microsoft YaHei", 14, "bold")).grid(row=0, column=0, columnspan=4, pady=(0, 15))
        
        row = 1
        ttk.Label(left_frame, text="姓    名：").grid(row=row, column=0, sticky="w", pady=5)
        self.name_entry = ttk.Entry(left_frame, width=20)
        self.name_entry.grid(row=row, column=1, sticky="w", pady=5)
        self.name_entry.focus_set()
        ttk.Label(left_frame, text="性    别：").grid(row=row, column=2, sticky="w", pady=5, padx=(20, 0))
        self.gender_var = tk.StringVar(value="男")
        ttk.Combobox(left_frame, textvariable=self.gender_var, values=["男", "女"], width=5, state="readonly").grid(row=row, column=3, sticky="w", pady=5)
        
        row += 1
        ttk.Label(left_frame, text="年    龄：").grid(row=row, column=0, sticky="w", pady=5)
        self.age_entry = ttk.Entry(left_frame, width=20)
        self.age_entry.grid(row=row, column=1, sticky="w", pady=5)
        ttk.Label(left_frame, text="电    话：").grid(row=row, column=2, sticky="w", pady=5, padx=(20, 0))
        self.phone_entry_patient = ttk.Entry(left_frame, width=20)
        self.phone_entry_patient.grid(row=row, column=3, sticky="w", pady=5)
        
        row += 1
        ttk.Label(left_frame, text="诊    断：").grid(row=row, column=0, sticky="w", pady=5)
        self.diagnosis_entry = ttk.Entry(left_frame, width=50)
        self.diagnosis_entry.grid(row=row, column=1, columnspan=3, sticky="w", pady=5)
        
        row += 1
        ttk.Label(left_frame, text="处    方：").grid(row=row, column=0, sticky="nw", pady=5)
        self.prescription_text = tk.Text(left_frame, width=50, height=6, undo=True)
        self.prescription_text.grid(row=row, column=1, columnspan=3, sticky="w", pady=5)
        scrollbar = ttk.Scrollbar(left_frame, orient="vertical", command=self.prescription_text.yview)
        scrollbar.grid(row=row, column=4, sticky="ns")
        self.prescription_text.configure(yscrollcommand=scrollbar.set)
        
        row += 1
        ttk.Label(left_frame, text="用    法：").grid(row=row, column=0, sticky="w", pady=5)
        self.usage_entry = ttk.Entry(left_frame, width=50)
        self.usage_entry.grid(row=row, column=1, columnspan=3, sticky="w", pady=5)
        self.usage_entry.insert(0, "水煎服，每日一剂，分早晚两次服用")
        
        row += 1
        ttk.Label(left_frame, text="开方医生：").grid(row=row, column=0, sticky="w", pady=5)
        self.doctor_label = ttk.Label(left_frame, text=self.settings.default_doctor or "未设置（请在设置页面配置）", foreground="blue")
        self.doctor_label.grid(row=row, column=1, sticky="w", pady=5)
        ttk.Label(left_frame, text=f"开方日期：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").grid(row=row, column=2, columnspan=2, sticky="w", pady=5, padx=(20, 0))
        
        row += 1
        ttk.Label(left_frame, text="打印机：").grid(row=row, column=0, sticky="w", pady=5)
        self.printer_combo = ttk.Combobox(left_frame, width=40, state="readonly")
        self.printer_combo.grid(row=row, column=1, columnspan=2, sticky="w", pady=5)
        ttk.Button(left_frame, text="刷新", command=self.refresh_printers, width=8).grid(row=row, column=3, sticky="w", pady=5, padx=(10, 0))
        
        row += 1
        btn_frame = ttk.Frame(left_frame)
        btn_frame.grid(row=row, column=0, columnspan=4, pady=15)
        ttk.Button(btn_frame, text="保存并打印", command=self.save_and_print, width=15).pack(side="left", padx=10)
        ttk.Button(btn_frame, text="仅保存", command=self.save_only, width=15).pack(side="left", padx=10)
        ttk.Button(btn_frame, text="清空表单", command=self.clear_form, width=15).pack(side="left", padx=10)
        
        self.root.bind('<Control-Return>', lambda e: self.save_and_print())
        self.root.bind('<Control-s>', lambda e: self.save_only())
        for entry in [self.name_entry, self.age_entry, self.phone_entry_patient, self.diagnosis_entry, self.usage_entry]:
            entry.bind('<Return>', lambda e: self.save_and_print())
        
        ContextMenu(self.name_entry)
        ContextMenu(self.age_entry)
        ContextMenu(self.phone_entry_patient)
        ContextMenu(self.diagnosis_entry)
        ContextMenu(self.prescription_text)
        ContextMenu(self.usage_entry)
        
        right_frame = ttk.Frame(main_container, padding="5")
        right_frame.pack(side="right", fill="both", padx=5, pady=5)
        
        print_frame = ttk.LabelFrame(right_frame, text="打印预览", padding="5")
        print_frame.pack(fill="both", expand=True)
        ttk.Label(print_frame, text="58mm小票格式", foreground="gray").pack(pady=(0, 2))
        self.print_preview_text = tk.Text(print_frame, width=35, height=25, bg="#f5f5f5", state="disabled")
        print_scroll = ttk.Scrollbar(print_frame, orient="vertical", command=self.print_preview_text.yview)
        self.print_preview_text.configure(yscrollcommand=print_scroll.set)
        print_scroll.pack(side="right", fill="y")
        self.print_preview_text.pack(side="left", fill="both", expand=True)
        
        for widget in [self.name_entry, self.age_entry, self.phone_entry_patient, self.diagnosis_entry, self.usage_entry]:
            widget.bind('<KeyRelease>', self.update_preview)
        self.prescription_text.bind('<KeyRelease>', self.update_preview)
        
        self.completion_panel = SmartCompletionPanel(main_container, self.db_file, self.settings, on_select_callback=self.insert_completion)
        self.completion_panel.get_frame().pack(fill="x", padx=10, pady=(0, 10))
        
        self._current_focused_widget = None
        for widget in [self.diagnosis_entry, self.prescription_text, self.usage_entry]:
            widget.bind('<FocusIn>', self._on_focus_change)
        
        self.update_preview()
    
    def _on_focus_change(self, event):
        self._current_focused_widget = event.widget
    
    def insert_completion(self, text):
        if self._current_focused_widget:
            widget = self._current_focused_widget
            if isinstance(widget, tk.Text):
                widget.insert(tk.INSERT, text)
                widget.focus_set()
            elif isinstance(widget, ttk.Entry):
                current = widget.get()
                pos = widget.index(tk.INSERT)
                new_text = current[:pos] + text + current[pos:]
                widget.delete(0, tk.END)
                widget.insert(0, new_text)
                widget.icursor(pos + len(text))
                widget.focus_set()
        self.update_preview()
    
    def create_query_page(self, parent):
        main_frame = ttk.Frame(parent, padding="20")
        main_frame.pack(fill="both", expand=True)
        ttk.Label(main_frame, text="历史处方查询", font=("Microsoft YaHei", 14, "bold")).pack(pady=(0, 20))
        
        # 搜索框架 - 包含姓名和日期搜索
        search_frame = ttk.LabelFrame(main_frame, text="查询条件", padding="10")
        search_frame.pack(fill="x", pady=10)
        
        # 第一行：姓名搜索
        name_frame = ttk.Frame(search_frame)
        name_frame.pack(fill="x", pady=5)
        ttk.Label(name_frame, text="患者姓名：", width=10).pack(side="left")
        self.search_entry = ttk.Entry(name_frame, width=30)
        self.search_entry.pack(side="left", padx=10)
        self.search_entry.bind('<Return>', lambda e: self.search_prescriptions())
        ContextMenu(self.search_entry)
        
        # 第二行：日期范围搜索
        date_frame = ttk.Frame(search_frame)
        date_frame.pack(fill="x", pady=5)
        ttk.Label(date_frame, text="日期范围：", width=10).pack(side="left")
        
        # 开始日期
        ttk.Label(date_frame, text="从：").pack(side="left", padx=(10, 5))
        self.start_date_entry = ttk.Entry(date_frame, width=15)
        self.start_date_entry.pack(side="left", padx=5)
        self.start_date_entry.insert(0, datetime.now().strftime("%Y-%m-01"))  # 默认本月第一天
        
        # 结束日期
        ttk.Label(date_frame, text="至：").pack(side="left", padx=(10, 5))
        self.end_date_entry = ttk.Entry(date_frame, width=15)
        self.end_date_entry.pack(side="left", padx=5)
        self.end_date_entry.insert(0, datetime.now().strftime("%Y-%m-%d"))  # 默认今天
        
        # 快速选择月份按钮
        ttk.Button(date_frame, text="本月", command=lambda: self.set_date_range("current_month"), width=8).pack(side="left", padx=10)
        ttk.Button(date_frame, text="上月", command=lambda: self.set_date_range("last_month"), width=8).pack(side="left", padx=5)
        ttk.Button(date_frame, text="近三月", command=lambda: self.set_date_range("recent_3_months"), width=8).pack(side="left", padx=5)
        
        # 查询按钮
        btn_frame = ttk.Frame(search_frame)
        btn_frame.pack(fill="x", pady=10)
        ttk.Button(btn_frame, text="查询", command=self.search_prescriptions, width=15).pack(side="left", padx=10)
        ttk.Button(btn_frame, text="全部/刷新", command=self.load_all_prescriptions, width=15).pack(side="left", padx=10)
        
        # 统计信息框架
        stats_frame = ttk.LabelFrame(main_frame, text="统计信息", padding="10")
        stats_frame.pack(fill="x", pady=10)
        
        self.stats_label = ttk.Label(stats_frame, text="总记录数：0 | 本月记录：0", font=("Microsoft YaHei", 10))
        self.stats_label.pack(side="left", padx=20)
        
        # 列表框架
        list_frame = ttk.LabelFrame(main_frame, text="处方记录列表", padding="10")
        list_frame.pack(fill="both", expand=True, pady=10)
        
        # 创建Treeview，添加更多列
        columns = ("ID", "姓名", "性别", "年龄", "日期", "中医辨证", "处方")
        self.tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=20)
        
        # 设置列标题和宽度
        self.tree.heading("ID", text="ID")
        self.tree.column("ID", width=50, anchor="center")
        self.tree.heading("姓名", text="姓名")
        self.tree.column("姓名", width=100, anchor="w")
        self.tree.heading("性别", text="性别")
        self.tree.column("性别", width=60, anchor="center")
        self.tree.heading("年龄", text="年龄")
        self.tree.column("年龄", width=80, anchor="center")
        self.tree.heading("日期", text="日期")
        self.tree.column("日期", width=120, anchor="center")
        self.tree.heading("中医辨证", text="中医辨证")
        self.tree.column("中医辨证", width=200, anchor="w")
        self.tree.heading("处方", text="处方")
        self.tree.column("处方", width=300, anchor="w")
        
        # 添加滚动条
        v_scroll = ttk.Scrollbar(list_frame, orient="vertical", command=self.tree.yview)
        h_scroll = ttk.Scrollbar(list_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=v_scroll.set, xscrollcommand=h_scroll.set)
        
        # 布局
        self.tree.grid(row=0, column=0, sticky="nsew")
        v_scroll.grid(row=0, column=1, sticky="ns")
        h_scroll.grid(row=1, column=0, sticky="ew")
        
        # 配置网格权重
        list_frame.grid_rowconfigure(0, weight=1)
        list_frame.grid_columnconfigure(0, weight=1)
        
        # 绑定双击事件
        self.tree.bind('<Double-1>', self.view_prescription_detail)
        
        # 操作按钮框架
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill="x", pady=10)
        ttk.Button(btn_frame, text="查看详情", command=self.view_prescription_detail, width=15).pack(side="left", padx=10)
        ttk.Button(btn_frame, text="打印此处方", command=self.print_selected_prescription, width=15).pack(side="left", padx=10)
        ttk.Button(btn_frame, text="删除此记录", command=self.delete_prescription, width=15).pack(side="left", padx=10)
        ttk.Button(btn_frame, text="导出数据", command=self.export_data, width=15).pack(side="left", padx=10)
        
        # 初始化加载数据
        self.load_all_prescriptions()
    
    def update_preview(self, event=None):
        self.update_print_preview()
    
    def update_print_preview(self):
        self.print_preview_text.config(state="normal")
        self.print_preview_text.delete("1.0", tk.END)
        preview = self.generate_receipt_text()
        self.print_preview_text.insert("1.0", preview)
        self.print_preview_text.config(state="disabled")
    
    def refresh_printers(self):
        try:
            printers = [p[2] for p in win32print.EnumPrinters(2)]
            self.printer_combo['values'] = printers
            if printers:
                for p in printers:
                    if any(keyword in p.lower() for keyword in ['pos', '58', 'receipt', 'thermal', '小票', '热敏']):
                        self.printer_combo.set(p)
                        return
                default_printer = win32print.GetDefaultPrinter()
                if default_printer in printers:
                    self.printer_combo.set(default_printer)
                else:
                    self.printer_combo.set(printers[0])
        except Exception as e:
            messagebox.showwarning("提示", f"获取打印机列表失败：{e}")
            self.printer_combo.set("")
    
    def save_only(self):
        if not self.validate_input():
            return
        if self.save_to_database():
            self.completion_panel.load_words_from_database()
            messagebox.showinfo("成功", "处方已保存！")
            self.clear_form()
    
    def save_and_print(self):
        if not self.validate_input():
            return
        if self.save_to_database():
            self.completion_panel.load_words_from_database()
            docx_file = self.generate_receipt_docx()
            if docx_file:
                if messagebox.askyesno("保存成功", f"处方已保存：\n{docx_file}\n\n是否立即打印？"):
                    self.print_docx(docx_file)
                messagebox.showinfo("完成", "处方已保存并打印！")
            else:
                messagebox.showinfo("成功", "处方已保存！")
            self.clear_form()
    
    def validate_input(self):
        name = self.name_entry.get().strip()
        prescription = self.prescription_text.get("1.0", "end-1c").strip()
        if not name:
            messagebox.showwarning("提示", "请输入患者姓名！")
            self.name_entry.focus_set()
            return False
        if not prescription:
            messagebox.showwarning("提示", "请输入处方内容！")
            self.prescription_text.focus_set()
            return False
        return True
    
    def save_to_database(self):
        try:
            conn = sqlite3.connect(self.db_file)
            cursor = conn.cursor()
            current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            cursor.execute("INSERT INTO prescriptions (patient_name, gender, age, phone, diagnosis, prescription, usage, doctor, doctor_phone, create_time) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (self.name_entry.get().strip(), self.gender_var.get(), self.age_entry.get().strip(),
                 self.phone_entry_patient.get().strip(), self.diagnosis_entry.get().strip(),
                 self.prescription_text.get("1.0", "end-1c").strip(), self.usage_entry.get().strip(),
                 self.settings.default_doctor, self.settings.default_phone, current_time))
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            messagebox.showerror("错误", f"保存失败：{e}")
            return False
    
    def generate_receipt_docx(self):
        """生成小票 - 使用用户配置的压缩参数"""
        if not DOCX_AVAILABLE:
            messagebox.showerror("错误", "未安装python-docx库！\n请运行: pip install python-docx")
            return None
        
        try:
            print(f"\n生成小票...")
            print(f"  字体大小: {self.settings.font_size}pt")
            print(f"  行间距: {self.settings.line_spacing}")
            print(f"  安全系数: {self.settings.safety_margin}")
            print(f"  边距: {self.settings.margin_size}cm")
            
            # 确保文件夹存在
            self.ensure_prescription_folder()
            
            doc = Document()
            section = doc.sections[0]
            
            # 设置页面宽度为58mm
            section.page_width = Cm(5.8)
            
            # 计算页面高度
            page_height = self.calculate_page_height()
            section.page_height = Cm(page_height)
            print(f"  计算页面高度: {page_height:.2f}cm")
            
            # 设置边距
            margin = self.settings.margin_size
            section.left_margin = Cm(margin)
            section.right_margin = Cm(margin)
            section.top_margin = Cm(margin)
            section.bottom_margin = Cm(margin)
            
            # 设置默认字体和段落格式
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(self.settings.font_size)
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            style.paragraph_format.line_spacing = self.settings.line_spacing
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            
            font_size = self.settings.font_size
            line_spacing = self.settings.line_spacing
            
            # 标题 - 使用设置的诊所名称
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.paragraph_format.line_spacing = line_spacing
            title.paragraph_format.space_before = Pt(0)
            title.paragraph_format.space_after = Pt(0)
            
            # 如果设置了诊所名称，使用诊所名称，否则使用默认值
            clinic_name = self.settings.clinic_name if self.settings.clinic_name else "海口市龙华区诊所名字"
            run = title.add_run(clinic_name)
            run.font.size = Pt(font_size + 2)
            run.font.bold = True
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            
            # 如果设置了诊所地址，添加地址信息
            if self.settings.clinic_address:
                address = doc.add_paragraph()
                address.alignment = WD_ALIGN_PARAGRAPH.CENTER
                address.paragraph_format.line_spacing = line_spacing
                address.paragraph_format.space_before = Pt(0)
                address.paragraph_format.space_after = Pt(0)
                address_run = address.add_run(self.settings.clinic_address)
                address_run.font.size = Pt(font_size)
                address_run.font.name = '宋体'
                address_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 如果设置了诊所电话，添加电话信息
            if self.settings.clinic_phone:
                phone = doc.add_paragraph()
                phone.alignment = WD_ALIGN_PARAGRAPH.CENTER
                phone.paragraph_format.line_spacing = line_spacing
                phone.paragraph_format.space_before = Pt(0)
                phone.paragraph_format.space_after = Pt(0)
                phone_run = phone.add_run(f"电话：{self.settings.clinic_phone}")
                phone_run.font.size = Pt(font_size)
                phone_run.font.name = '宋体'
                phone_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 副标题
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.paragraph_format.line_spacing = line_spacing
            subtitle.paragraph_format.space_before = Pt(0)
            subtitle.paragraph_format.space_after = Pt(0)
            run = subtitle.add_run("中医干预中药处方")
            run.font.size = Pt(font_size + 1)
            run.font.bold = True
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            
            # 分隔线
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = line_spacing
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run("─" * 16)

            run.font.size = Pt(font_size)
            
            # 患者信息
            self._add_compact_line(doc, f"姓名：{self.name_entry.get().strip()}", font_size, line_spacing)
            self._add_compact_line(doc, f"性别：{self.gender_var.get()}  年龄：{self.age_entry.get().strip()}", font_size, line_spacing)
            self._add_compact_line(doc, f"电话：{self.phone_entry_patient.get().strip()}", font_size, line_spacing)
            self._add_compact_line(doc, f"日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}", font_size, line_spacing)
            
            # 分隔线
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = line_spacing
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run("─" * 16)
            run.font.size = Pt(font_size)
            
            # 中医辨证
            diagnosis = self.diagnosis_entry.get().strip()
            if diagnosis:
                para = doc.add_paragraph()
                para.paragraph_format.line_spacing = line_spacing
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                run = para.add_run("中医辨证：")
                run.font.bold = True
                run.font.size = Pt(font_size)
                run = para.add_run(diagnosis)
                run.font.size = Pt(font_size)
            
            # 处方
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = line_spacing
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run("处方：")
            run.font.bold = True
            run.font.size = Pt(font_size)
            
            prescription_content = self.prescription_text.get("1.0", "end-1c").strip()
            if prescription_content:
                prescription_lines = prescription_content.split('\n')
                # 限制处方最大行数，确保单页
                max_prescription_lines = 15  # 限制处方最多15行
                line_count = 0
                for line in prescription_lines:
                    if line.strip():
                        line_count += 1
                        if line_count > max_prescription_lines:
                            # 添加省略号表示内容被截断
                            self._add_compact_line(doc, "  ...（内容过多，已截断）", font_size, line_spacing)
                            break
                        self._add_compact_line(doc, f"  {line.strip()}", font_size, line_spacing)
            
            # 用法
            usage = self.usage_entry.get().strip()
            if usage:
                para = doc.add_paragraph()
                para.paragraph_format.line_spacing = line_spacing
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                run = para.add_run("用法：")
                run.font.bold = True
                run.font.size = Pt(font_size)
                run = para.add_run(usage)
                run.font.size = Pt(font_size)
            
            # 分隔线
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = line_spacing
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run("=" * 34)
            run.font.size = Pt(font_size)
            
            # 医生信息
            doctor = self.settings.default_doctor
            if doctor:
                self._add_compact_line(doc, f"开方医生：{doctor}", font_size, line_spacing)
            phone = self.settings.default_phone
            if phone:
                self._add_compact_line(doc, f"联系电话：{phone}", font_size, line_spacing)
            
            # 添加执业许可证号
            license_num = self.settings.clinic_license
            if license_num:
                self._add_compact_line(doc, f"执业许可证号：{license_num}", font_size, line_spacing)
            
            # 最后分隔线
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = line_spacing
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run("=" * 34)
            run.font.size = Pt(font_size)
            
            # 生成文件名并保存
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            patient_name = self.name_entry.get().strip()
            filename = os.path.join(PRESCRIPTION_FOLDER, f"处方_{patient_name}_{timestamp}.docx")
            doc.save(filename)
            
            print(f"  ✓ 小票生成成功！")
            
            # 保存txt文件
            try:
                txt_file = filename.replace('.docx', '.txt')
                with open(txt_file, "w", encoding="utf-8") as f:
                    f.write(self.generate_receipt_text())
            except:
                pass
            
            return os.path.abspath(filename)
            
        except Exception as e:
            print(f"  生成失败: {e}")
            messagebox.showerror("错误", f"生成小票失败：{e}")
            return None
    
    def _add_compact_line(self, doc, text, font_size, line_spacing):
        """添加紧凑行"""
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = line_spacing
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(text)
        run.font.size = Pt(font_size)
    
    def calculate_page_height(self):
        """计算页面高度 - 强制单页"""
        # 基础行高
        base_line_height = (self.settings.font_size / 72) * 2.54 * self.settings.line_spacing
        
        # 计算内容总高度
        content_height = 0
        
        # 标题区域
        content_height += base_line_height * 1.2  # 医院名称
        
        # 新增：诊所地址（如果有）
        if self.settings.clinic_address:
            address_lines = max(1, len(self.settings.clinic_address) // 18 + 1)
            content_height += base_line_height * address_lines
            
        # 新增：诊所电话（如果有）
        if self.settings.clinic_phone:
            content_height += base_line_height
            
        content_height += base_line_height * 1.0  # 处方标题
        content_height += base_line_height * 0.6  # 分隔线
        
        # 患者信息（4行）
        content_height += base_line_height * 4
        
        # 分隔线
        content_height += base_line_height * 0.6
        
        # 中医辨证
        diagnosis = self.diagnosis_entry.get().strip()
        if diagnosis:
            diag_lines = max(1, len(diagnosis) // 18 + 1)
            content_height += base_line_height * diag_lines
        
        # 处方标签
        content_height += base_line_height
        
        # 处方内容 - 限制最大行数
        prescription_content = self.prescription_text.get("1.0", "end-1c").strip()
        if prescription_content:
            prescription_lines = prescription_content.split('\n')
            # 限制处方最大行数，确保单页
            max_prescription_lines = 15  # 限制处方最多15行
            line_count = 0
            for line in prescription_lines:
                if line.strip():
                    line_count += 1
                    if line_count > max_prescription_lines:
                        break
                    line_lines = max(1, len(line.strip()) // 16 + 1)
                    content_height += base_line_height * line_lines
        
        # 用法
        usage = self.usage_entry.get().strip()
        if usage:
            usage_lines = max(1, len(usage) // 18 + 1)
            content_height += base_line_height * usage_lines
        
        # 分隔线
        content_height += base_line_height * 0.6
        
        # 医生信息
        if self.settings.default_doctor:
            content_height += base_line_height
        if self.settings.default_phone:
            content_height += base_line_height
            
        # 新增：执业许可证号（如果有）
        if self.settings.clinic_license:
            content_height += base_line_height
        
        # 最后分隔线
        content_height += base_line_height * 0.6
        
        # 计算总高度，使用留空比例控制
        # 留空比例 = 总高度 / 内容高度，例如4/1表示总高度是内容高度的4倍
        empty_ratio = self.settings.empty_ratio
        total_height = content_height * empty_ratio
        
        # 添加边距
        total_height += self.settings.margin_size * 2  # 上下边距
        
        # 严格限制最大高度为29.7cm（A4纸高度），确保单页
        max_height = 29.7
        if total_height > max_height:
            total_height = max_height
            
        return total_height
    
    def generate_receipt_text(self):
        lines = []
        lines.append("=" * 22)
        
        # 使用设置中的诊所名称，如果未设置则使用默认值
        clinic_name = self.settings.clinic_name if self.settings.clinic_name else "海口市龙华区诊所名字"
        lines.append(f" {clinic_name}")
        
        lines.append("   中医干预中药处方")
        lines.append("=" * 22)
        lines.append(f"姓名：{self.name_entry.get().strip()}")
        lines.append(f"性别：{self.gender_var.get()}  年龄：{self.age_entry.get().strip()}")
        lines.append(f"电话：{self.phone_entry_patient.get().strip()}")
        lines.append(f"日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        lines.append("-" * 22)
        diagnosis = self.diagnosis_entry.get().strip()
        if diagnosis:
            lines.append(f"中医辨证：{diagnosis}")
        lines.append("")
        lines.append("处方：")
        prescription_content = self.prescription_text.get("1.0", "end-1c").strip()
        if prescription_content:
            prescription_lines = prescription_content.split('\n')
            # 限制处方最大行数，确保单页
            max_prescription_lines = 15  # 限制处方最多15行
            line_count = 0
            for line in prescription_lines:
                if line.strip():
                    line_count += 1
                    if line_count > max_prescription_lines:
                        lines.append("  ...（内容过多，已截断）")
                        break
                    lines.append(f"  {line.strip()}")
        lines.append("")
        usage = self.usage_entry.get().strip()
        if usage:
            lines.append(f"用法：{usage}")
        lines.append("-" * 22)
        doctor = self.settings.default_doctor
        if doctor:
            lines.append(f"开方医生：{doctor}")
        phone = self.settings.default_phone
        if phone:
            lines.append(f"联系电话：{phone}")
        
        # 添加执业许可证号
        license_num = self.settings.clinic_license
        if license_num:
            lines.append(f"执业许可证号：{license_num}")
        lines.append("=" * 22)
        return '\n'.join(lines)
    
    def print_docx(self, docx_file):
        printer_name = self.printer_combo.get()
        try:
            if printer_name:
                # 尝试多种打印方法
                success = False
                
                # 方法1：使用os.startfile直接打印到指定打印机
                try:
                    # 使用win32print设置默认打印机，然后使用os.startfile
                    import win32print
                    import win32api
                    
                    # 保存当前默认打印机
                    current_printer = win32print.GetDefaultPrinter()
                    
                    # 临时设置默认打印机
                    try:
                        win32print.SetDefaultPrinter(printer_name)
                        # 打印文件
                        os.startfile(docx_file, "print")
                        success = True
                    finally:
                        # 恢复原来的默认打印机
                        if current_printer:
                            win32print.SetDefaultPrinter(current_printer)
                except Exception as e1:
                    print(f"默认打印机切换方法失败: {e1}")
                
                # 方法2：如果方法1失败，尝试使用PowerShell简单打印
                if not success:
                    try:
                        import subprocess
                        # 使用简单的PowerShell命令
                        cmd = f'& Start-Process -FilePath "{docx_file}" -Verb Print'
                        subprocess.run(['powershell', '-Command', cmd], 
                                     check=True, timeout=30, 
                                     creationflags=subprocess.CREATE_NO_WINDOW)
                        success = True
                    except Exception as e2:
                        print(f"PowerShell简单打印失败: {e2}")
                
                # 方法3：如果以上都失败，尝试直接使用win32api
                if not success:
                    try:
                        import win32api
                        # 使用print命令
                        win32api.ShellExecute(0, "print", docx_file, "", ".", 0)
                        success = True
                    except Exception as e3:
                        print(f"win32api打印失败: {e3}")
                
                # 如果所有方法都失败，显示错误并提供备选方案
                if not success:
                    raise Exception("所有打印方法都失败，请检查打印机连接和驱动程序")
            else:
                # 如果没有指定打印机，使用默认方式
                os.startfile(docx_file, "print")
                
        except Exception as e:
            messagebox.showerror("打印错误", f"打印失败：{e}\n\n可能的解决方案：\n1. 检查打印机是否已连接并开机\n2. 检查打印机驱动程序是否正常\n3. 检查打印机是否在线\n4. 尝试以管理员身份运行程序")
            try:
                # 最后的备选方案：直接打开文件
                os.startfile(docx_file)
            except:
                pass
    
    def clear_form(self):
        self.name_entry.delete(0, tk.END)
        self.gender_var.set("男")
        self.age_entry.delete(0, tk.END)
        self.phone_entry_patient.delete(0, tk.END)
        self.diagnosis_entry.delete(0, tk.END)
        self.prescription_text.delete("1.0", tk.END)
        self.usage_entry.delete(0, tk.END)
        self.usage_entry.insert(0, "水煎服，每日一剂，分早晚两次服用")
        if self.settings.default_doctor:
            self.doctor_label.config(text=self.settings.default_doctor)
        else:
            self.doctor_label.config(text="未设置（请在设置页面配置）")
        self.name_entry.focus_set()
        self.update_preview()
    
    def set_date_range(self, range_type):
        """设置日期范围"""
        now = datetime.now()
        if range_type == "current_month":
            # 本月第一天到今天
            start_date = now.replace(day=1).strftime("%Y-%m-%d")
            end_date = now.strftime("%Y-%m-%d")
        elif range_type == "last_month":
            # 上月第一天到上月最后一天
            if now.month == 1:
                last_month = now.replace(year=now.year-1, month=12, day=1)
            else:
                last_month = now.replace(month=now.month-1, day=1)
            
            # 计算上月最后一天
            if last_month.month == 12:
                next_month = last_month.replace(year=last_month.year+1, month=1)
            else:
                next_month = last_month.replace(month=last_month.month+1)
            
            last_day_of_last_month = (next_month - timedelta(days=1)).day
            end_of_last_month = last_month.replace(day=last_day_of_last_month)
            
            start_date = last_month.strftime("%Y-%m-%d")
            end_date = end_of_last_month.strftime("%Y-%m-%d")
        elif range_type == "recent_3_months":
            # 近三个月
            if now.month > 3:
                three_months_ago = now.replace(month=now.month-3)
            else:
                three_months_ago = now.replace(year=now.year-1, month=now.month+9)
            start_date = three_months_ago.strftime("%Y-%m-%d")
            end_date = now.strftime("%Y-%m-%d")
        else:
            return
        
        # 更新输入框
        self.start_date_entry.delete(0, tk.END)
        self.start_date_entry.insert(0, start_date)
        self.end_date_entry.delete(0, tk.END)
        self.end_date_entry.insert(0, end_date)
        
        # 自动执行查询
        self.search_prescriptions()

    def search_prescriptions(self):
        search_name = self.search_entry.get().strip()
        start_date = self.start_date_entry.get().strip()
        end_date = self.end_date_entry.get().strip()
        now = datetime.now()
        
        try:
            conn = sqlite3.connect(self.db_file)
            cursor = conn.cursor()
            
            # 构建查询条件
            conditions = []
            params = []
            
            if search_name:
                conditions.append("patient_name LIKE ?")
                params.append(f'%{search_name}%')
            
            if start_date:
                conditions.append("DATE(create_time) >= ?")
                params.append(start_date)
            
            if end_date:
                conditions.append("DATE(create_time) <= ?")
                params.append(end_date)
            
            # 构建SQL查询
            if conditions:
                query = f"SELECT id, patient_name, gender, age, create_time, diagnosis, prescription FROM prescriptions WHERE {' AND '.join(conditions)} ORDER BY create_time DESC"
            else:
                query = "SELECT id, patient_name, gender, age, create_time, diagnosis, prescription FROM prescriptions ORDER BY create_time DESC"
            
            cursor.execute(query, params)
            rows = cursor.fetchall()
            
            # 获取统计信息
            cursor.execute("SELECT COUNT(*) FROM prescriptions")
            total_count = cursor.fetchone()[0]
            
            # 获取本月记录数
            current_month_start = now.replace(day=1).strftime("%Y-%m-%d")
            cursor.execute("SELECT COUNT(*) FROM prescriptions WHERE DATE(create_time) >= ?", (current_month_start,))
            month_count = cursor.fetchone()[0]
            
            conn.close()
            
            # 更新统计信息
            self.stats_label.config(text=f"总记录数：{total_count} | 本月记录：{month_count} | 当前显示：{len(rows)}")
            
            # 清空并重新填充表格
            for item in self.tree.get_children():
                self.tree.delete(item)
            
            for row in rows:
                prescription_short = row[6][:50] + "..." if len(row[6]) > 50 else row[6]
                # 格式化日期显示
                date_str = row[4][:10] if row[4] else ""
                self.tree.insert("", "end", values=(row[0], row[1], row[2] or "", row[3] or "", date_str, row[5] or "", prescription_short), tags=(row[0],))
                
        except Exception as e:
            messagebox.showerror("错误", f"查询失败：{e}")
    
    def load_all_prescriptions(self):
        self.search_entry.delete(0, tk.END)
        self.search_prescriptions()
    
    def view_prescription_detail(self, event=None):
        selection = self.tree.selection()
        if not selection:
            messagebox.showwarning("提示", "请先选择一条记录！")
            return
        item = selection[0]
        prescription_id = self.tree.item(item, "tags")[0]
        try:
            conn = sqlite3.connect(self.db_file)
            cursor = conn.cursor()
            cursor.execute("SELECT patient_name, gender, age, phone, diagnosis, prescription, usage, doctor, doctor_phone, create_time FROM prescriptions WHERE id = ?", (prescription_id,))
            row = cursor.fetchone()
            conn.close()
            if row:
                # 创建现代化的详情窗口
                detail_window = tk.Toplevel(self.root)
                detail_window.title("处方详情")
                detail_window.geometry("700x650")
                detail_window.configure(bg="#f0f0f0")
                
                # 创建主框架
                main_frame = tk.Frame(detail_window, bg="#f0f0f0", padx=20, pady=20)
                main_frame.pack(fill="both", expand=True)
                
                # 标题
                title_label = tk.Label(main_frame, text="处方详情", font=("Microsoft YaHei", 18, "bold"), fg="#2c3e50")
                title_label.pack(pady=(0, 20))
                
                # 患者信息框架
                patient_frame = tk.Frame(main_frame, bg="#f0f0f0", relief="groove", bd=2)
                patient_frame.pack(fill="x", pady=10)
                # 添加患者信息标签
                patient_label = tk.Label(patient_frame, text="  患者信息  ", font=("Microsoft YaHei", 10, "bold"), bg="#f0f0f0", fg="#2c3e50")
                patient_label.pack(anchor="w", padx=5)
                # 添加内容框架
                patient_content = tk.Frame(patient_frame, bg="#f0f0f0", padx=15, pady=5)
                patient_content.pack(fill="x")
                
                # 患者信息 - 使用不同颜色
                info_frame = tk.Frame(patient_content, bg="#f0f0f0")
                info_frame.pack(fill="x")
                
                # 姓名 - 蓝色
                name_frame = tk.Frame(info_frame, bg="#f0f0f0")
                name_frame.pack(fill="x", pady=5)
                tk.Label(name_frame, text="姓名：", font=("Microsoft YaHei", 12, "bold"), fg="#3498db", width=10, anchor="w").pack(side="left")
                tk.Label(name_frame, text=row[0], font=("Microsoft YaHei", 12), fg="#2c3e50", anchor="w").pack(side="left")
                
                # 性别和年龄 - 绿色
                gender_age_frame = tk.Frame(info_frame, bg="#f0f0f0")
                gender_age_frame.pack(fill="x", pady=5)
                tk.Label(gender_age_frame, text="性别：", font=("Microsoft YaHei", 12, "bold"), fg="#27ae60", width=10, anchor="w").pack(side="left")
                tk.Label(gender_age_frame, text=row[1] or "未填写", font=("Microsoft YaHei", 12), fg="#2c3e50", anchor="w").pack(side="left")
                tk.Label(gender_age_frame, text="  年龄：", font=("Microsoft YaHei", 12, "bold"), fg="#27ae60", anchor="w").pack(side="left")
                tk.Label(gender_age_frame, text=row[2] or "未填写", font=("Microsoft YaHei", 12), fg="#2c3e50", anchor="w").pack(side="left")
                
                # 电话 - 紫色
                phone_frame = tk.Frame(info_frame, bg="#f0f0f0")
                phone_frame.pack(fill="x", pady=5)
                tk.Label(phone_frame, text="电话：", font=("Microsoft YaHei", 12, "bold"), fg="#9b59b6", width=10, anchor="w").pack(side="left")
                tk.Label(phone_frame, text=row[3] or "未填写", font=("Microsoft YaHei", 12), fg="#2c3e50", anchor="w").pack(side="left")
                
                # 就诊日期 - 橙色
                date_frame = tk.Frame(info_frame, bg="#f0f0f0")
                date_frame.pack(fill="x", pady=5)
                tk.Label(date_frame, text="就诊日期：", font=("Microsoft YaHei", 12, "bold"), fg="#e67e22", width=10, anchor="w").pack(side="left")
                tk.Label(date_frame, text=row[9], font=("Microsoft YaHei", 12), fg="#2c3e50", anchor="w").pack(side="left")
                
                # 诊断信息框架
                diagnosis_frame = tk.Frame(main_frame, bg="#f0f0f0", relief="groove", bd=2)
                diagnosis_frame.pack(fill="x", pady=10)
                # 添加诊断信息标签
                diagnosis_label = tk.Label(diagnosis_frame, text="  诊断信息  ", font=("Microsoft YaHei", 10, "bold"), bg="#f0f0f0", fg="#2c3e50")
                diagnosis_label.pack(anchor="w", padx=5)
                # 添加内容框架
                diagnosis_content = tk.Frame(diagnosis_frame, bg="#f0f0f0", padx=15, pady=5)
                diagnosis_content.pack(fill="x")
                
                # 中医辨证 - 红色
                diagnosis_text_frame = tk.Frame(diagnosis_content, bg="#f0f0f0")
                diagnosis_text_frame.pack(fill="x")
                tk.Label(diagnosis_text_frame, text="中医辨证：", font=("Microsoft YaHei", 12, "bold"), fg="#e74c3c", width=10, anchor="nw").pack(side="left")
                diagnosis_label = tk.Label(diagnosis_text_frame, text=row[4] or "未填写", font=("Microsoft YaHei", 12), fg="#2c3e50", justify="left", anchor="nw")
                diagnosis_label.pack(side="left", fill="x", expand=True)
                
                # 处方信息框架
                prescription_frame = tk.Frame(main_frame, bg="#f0f0f0", relief="groove", bd=2)
                prescription_frame.pack(fill="both", expand=True, pady=10)
                # 添加处方信息标签
                prescription_label = tk.Label(prescription_frame, text="  处方信息  ", font=("Microsoft YaHei", 10, "bold"), bg="#f0f0f0", fg="#2c3e50")
                prescription_label.pack(anchor="w", padx=5)
                # 添加内容框架
                prescription_content = tk.Frame(prescription_frame, bg="#f0f0f0", padx=15, pady=5)
                prescription_content.pack(fill="both", expand=True)
                
                # 处方内容 - 深蓝色
                prescription_text_frame = tk.Frame(prescription_content, bg="#f0f0f0")
                prescription_text_frame.pack(fill="both", expand=True)
                tk.Label(prescription_text_frame, text="处方内容：", font=("Microsoft YaHei", 12, "bold"), fg="#2980b9", width=10, anchor="nw").pack(side="left")
                
                # 处方内容文本框
                prescription_text = tk.Text(prescription_text_frame, height=8, width=50, font=("Microsoft YaHei", 11), wrap=tk.WORD, relief="solid", borderwidth=1, bg="#f0f0f0")
                prescription_text.pack(side="left", fill="both", expand=True, padx=(10, 0))
                prescription_text.insert("1.0", row[5] or "")
                prescription_text.config(state="disabled")
                
                # 用法信息 - 棕色
                usage_frame = tk.Frame(prescription_content, bg="#f0f0f0")
                usage_frame.pack(fill="x", pady=(10, 0))
                tk.Label(usage_frame, text="用法：", font=("Microsoft YaHei", 12, "bold"), fg="#8e44ad", width=10, anchor="w").pack(side="left")
                tk.Label(usage_frame, text=row[6] or "未填写", font=("Microsoft YaHei", 12), fg="#2c3e50", anchor="w").pack(side="left")
                
                # 医生信息框架
                doctor_frame = tk.Frame(main_frame, bg="#f0f0f0", relief="groove", bd=2)
                doctor_frame.pack(fill="x", pady=10)
                # 添加医生信息标签
                doctor_label = tk.Label(doctor_frame, text="  医生信息  ", font=("Microsoft YaHei", 10, "bold"), bg="#f0f0f0", fg="#2c3e50")
                doctor_label.pack(anchor="w", padx=5)
                # 添加内容框架
                doctor_content = tk.Frame(doctor_frame, bg="#f0f0f0", padx=15, pady=5)
                doctor_content.pack(fill="x")
                
                # 开方医生 - 青色
                doctor_info_frame = tk.Frame(doctor_content, bg="#f0f0f0")
                doctor_info_frame.pack(fill="x")
                tk.Label(doctor_info_frame, text="开方医生：", font=("Microsoft YaHei", 12, "bold"), fg="#16a085", width=10, anchor="w").pack(side="left")
                tk.Label(doctor_info_frame, text=row[7] or "未填写", font=("Microsoft YaHei", 12), fg="#2c3e50", anchor="w").pack(side="left")
                
                # 医生电话 - 深绿色
                doctor_phone_frame = tk.Frame(doctor_content, bg="#f0f0f0")
                doctor_phone_frame.pack(fill="x", pady=5)
                tk.Label(doctor_phone_frame, text="医生电话：", font=("Microsoft YaHei", 12, "bold"), fg="#27ae60", width=10, anchor="w").pack(side="left")
                tk.Label(doctor_phone_frame, text=row[8] or "未填写", font=("Microsoft YaHei", 12), fg="#2c3e50", anchor="w").pack(side="left")
                
                # 按钮框架
                btn_frame = tk.Frame(main_frame, bg="#f0f0f0")
                btn_frame.pack(fill="x", pady=20)
                
                # 打印按钮
                print_btn = ttk.Button(btn_frame, text="打印此处方", command=lambda: self.print_prescription_by_id(prescription_id))
                print_btn.pack(side="left", padx=10)
                
                # 关闭按钮
                close_btn = ttk.Button(btn_frame, text="关闭", command=detail_window.destroy)
                close_btn.pack(side="right", padx=10)
                
        except Exception as e:
            messagebox.showerror("错误", f"获取详情失败：{e}")
    
    def print_prescription_by_id(self, prescription_id):
        """根据ID打印处方"""
        try:
            conn = sqlite3.connect(self.db_file)
            cursor = conn.cursor()
            cursor.execute("SELECT patient_name, gender, age, phone, diagnosis, prescription, usage, doctor, doctor_phone, create_time FROM prescriptions WHERE id = ?", (prescription_id,))
            row = cursor.fetchone()
            conn.close()
            if row:
                # 填充表单
                self.name_entry.delete(0, tk.END)
                self.name_entry.insert(0, row[0])
                self.gender_var.set(row[1] or "男")
                self.age_entry.delete(0, tk.END)
                self.age_entry.insert(0, row[2] or "")
                self.phone_entry_patient.delete(0, tk.END)
                self.phone_entry_patient.insert(0, row[3] or "")
                self.diagnosis_entry.delete(0, tk.END)
                self.diagnosis_entry.insert(0, row[4] or "")
                self.prescription_text.delete("1.0", tk.END)
                self.prescription_text.insert("1.0", row[5] or "")
                self.usage_entry.delete(0, tk.END)
                self.usage_entry.insert(0, row[6] or "水煎服，每日一剂，分早晚两次服用")
                
                # 更新预览
                self.update_preview()
                
                # 生成并打印文档
                docx_file = self.generate_receipt_docx()
                if docx_file:
                    if messagebox.askyesno("打印", f"处方已生成：\n{docx_file}\n\n是否打印？"):
                        self.print_docx(docx_file)
        except Exception as e:
            messagebox.showerror("错误", f"打印失败：{e}")

    def print_selected_prescription(self):
        selection = self.tree.selection()
        if not selection:
            messagebox.showwarning("提示", "请先选择一条记录！")
            return
        item = selection[0]
        prescription_id = self.tree.item(item, "tags")[0]
        try:
            conn = sqlite3.connect(self.db_file)
            cursor = conn.cursor()
            cursor.execute("SELECT patient_name, gender, age, phone, diagnosis, prescription, usage, doctor, doctor_phone, create_time FROM prescriptions WHERE id = ?", (prescription_id,))
            row = cursor.fetchone()
            conn.close()
            if row:
                self.name_entry.delete(0, tk.END)
                self.name_entry.insert(0, row[0])
                self.gender_var.set(row[1] or "男")
                self.age_entry.delete(0, tk.END)
                self.age_entry.insert(0, row[2] or "")
                self.phone_entry_patient.delete(0, tk.END)
                self.phone_entry_patient.insert(0, row[3] or "")
                self.diagnosis_entry.delete(0, tk.END)
                self.diagnosis_entry.insert(0, row[4] or "")
                self.prescription_text.delete("1.0", tk.END)
                self.prescription_text.insert("1.0", row[5] or "")
                self.usage_entry.delete(0, tk.END)
                self.usage_entry.insert(0, row[6] or "水煎服，每日一剂，分早晚两次服用")
                self.update_preview()
                docx_file = self.generate_receipt_docx()
                if docx_file:
                    if messagebox.askyesno("打印", f"处方已生成：\n{docx_file}\n\n是否打印？"):
                        self.print_docx(docx_file)
        except Exception as e:
            messagebox.showerror("错误", f"打印失败：{e}")
    
    def delete_prescription(self):
        selection = self.tree.selection()
        if not selection:
            messagebox.showwarning("提示", "请先选择一条记录！")
            return
        if not messagebox.askyesno("确认", "确定要删除这条记录吗？"):
            return
        item = selection[0]
        prescription_id = self.tree.item(item, "tags")[0]
        try:
            conn = sqlite3.connect(self.db_file)
            cursor = conn.cursor()
            cursor.execute("DELETE FROM prescriptions WHERE id = ?", (prescription_id,))
            conn.commit()
            conn.close()
            self.completion_panel.load_words_from_database()
            messagebox.showinfo("成功", "记录已删除！")
            self.search_prescriptions()
        except Exception as e:
            messagebox.showerror("错误", f"删除失败：{e}")
    
    def export_data(self):
        try:
            conn = sqlite3.connect(self.db_file)
            cursor = conn.cursor()
            cursor.execute("SELECT patient_name, gender, age, phone, diagnosis, prescription, usage, doctor, doctor_phone, create_time FROM prescriptions ORDER BY create_time DESC")
            rows = cursor.fetchall()
            conn.close()
            if not rows:
                messagebox.showinfo("提示", "没有数据可导出！")
                return
            filename = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV文件", "*.csv"), ("所有文件", "*.*")], title="导出数据")
            if not filename:
                return
            with open(filename, "w", encoding="utf-8-sig") as f:
                f.write("姓名,性别,年龄,电话,中医辨证,处方,用法,医生,医生电话,日期\n")
                for row in rows:
                    line = ",".join([f'"{str(field)}"' for field in row])
                    f.write(line + "\n")
            messagebox.showinfo("成功", f"数据已导出到：\n{filename}")
        except Exception as e:
            messagebox.showerror("错误", f"导出失败：{e}")


def main():
    root = tk.Tk()
    style = ttk.Style()
    style.theme_use('clam')
    app = PrescriptionSystem(root)
    root.update_idletasks()
    width = root.winfo_width()
    height = root.winfo_height()
    x = (root.winfo_screenwidth() // 2) - (width // 2)
    y = (root.winfo_screenheight() // 2) - (height // 2)
    root.geometry(f'{width}x{height}+{x}+{y}')
    root.mainloop()


if __name__ == "__main__":
    main()